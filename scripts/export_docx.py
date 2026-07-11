"""Xuất báo cáo demo ra Word (.docx) kèm sơ đồ -> BAO_CAO_DEMO.docx.

Dựng trực tiếp bằng python-docx (không cần pandoc). Ảnh lấy từ reports/figures/.
Chạy:  python scripts/export_docx.py
"""
import _bootstrap  # noqa: F401

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

ROOT = _bootstrap.ROOT
FIG = _bootstrap.REPORTS / "figures"
OUT = ROOT / "BAO_CAO_DEMO.docx"

GREY = RGBColor(0x55, 0x55, 0x55)


def main():
    doc = Document()

    # Font mặc định (Calibri có đủ tiếng Việt).
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    def img(name, width=15.5):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(FIG / name), width=Cm(width))

    def caption(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = GREY

    def note(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.italic = True
        r.font.color.rgb = GREY

    def table(headers, rows, widths=None):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Light Grid Accent 1"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for c, h in zip(t.rows[0].cells, headers):
            c.paragraphs[0].add_run(h).bold = True
        for row in rows:
            cells = t.add_row().cells
            for c, val in zip(cells, row):
                c.text = str(val)
        if widths:
            for col, w in zip(t.columns, widths):
                for c in col.cells:
                    c.width = Cm(w)
        doc.add_paragraph()

    # ---------------------------------------------------------------- TITLE
    h = doc.add_heading("BÁO CÁO DEMO — WiseTravel", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Giải quyết bài toán dữ liệu cho ứng dụng du lịch")
    r.bold = True
    r.font.size = Pt(13)
    for line in ("Loại: Demo kỹ thuật (không phải sản phẩm hoàn chỉnh).",
                 "Phạm vi: 12 quận nội thành Hà Nội · ~500 địa điểm.",
                 "Cập nhật: 11/07/2026."):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = p.add_run(line)
        rr.font.size = Pt(10)
        rr.font.color.rgb = GREY

    # ---------------------------------------------------------------- SUMMARY
    doc.add_heading("Tóm tắt kết quả", level=1)
    table(
        ["Hạng mục", "Kết quả", "Ý nghĩa"],
        [["Địa điểm thu thập", "500 POI thật (OpenStreetMap)", "Dữ liệu thật, không bịa"],
         ["Audit — tọa độ", "64.0% (32/50)", "Tự đo được chất lượng"],
         ["Audit — giờ mở cửa", "47.9% (23/48)", "Phát hiện giới hạn dữ liệu mở"],
         ["Cohen's Kappa", "0.611 (substantial)", "Tập gold đáng tin"],
         ["Model tốt nhất", "ViSoBERT — acc 86.4%, F1 0.723", "Chọn model bằng bằng chứng"],
         ["Sản phẩm demo", "App Streamlit 3 tab", "Gói toàn bộ bằng chứng"]],
        widths=[4.5, 6.0, 5.0],
    )

    # ---------------------------------------------------------------- 1
    doc.add_heading("1. Bài toán & mục tiêu", level=1)
    doc.add_paragraph(
        "Một ứng dụng du lịch cần hai loại dữ liệu cốt lõi: danh sách địa điểm (ở đâu, mở cửa "
        "khi nào, giá tầm nào) và cảm nhận của người dùng (được khen hay chê). Cả hai đều là "
        "“bài toán dữ liệu”: thu thập thì dễ, nhưng biết dữ liệu đúng hay sai mới khó.")
    doc.add_paragraph(
        "Đồ án chứng minh nhóm làm được quy trình dữ liệu có kiểm chứng định lượng cho cả hai "
        "loại, thay vì thu thập rồi tin tưởng cảm tính.")

    # ---------------------------------------------------------------- 2
    doc.add_heading("2. Kiến trúc & quy trình", level=1)
    doc.add_paragraph("Hai nhánh dữ liệu độc lập, hội tụ vào một app trình diễn:")
    img("fig_pipeline.png", 16)
    caption("Sơ đồ quy trình hai nhánh dữ liệu")
    doc.add_paragraph(
        "Công nghệ: Python · SQLite · OpenStreetMap (Overpass, miễn phí) · model HuggingFace "
        "(không tự huấn luyện) · Streamlit. Toàn bộ chạy local, tự chứa, không cần dịch vụ trả phí.")

    # ---------------------------------------------------------------- 3
    doc.add_heading("3. Nhánh 1 — Dữ liệu địa điểm", level=1)

    doc.add_heading("3.1 Thu thập & chuẩn hóa", level=2)
    doc.add_paragraph(
        "Lấy địa điểm thật từ OpenStreetMap qua Overpass API, gom về 4 loại hình, khử trùng lặp, "
        "lưu vào SQLite. Kết quả 500 địa điểm:")
    img("fig_category.png", 13)
    caption("Phân bố 500 địa điểm theo loại hình")

    doc.add_heading("3.2 Minh bạch nguồn giờ mở cửa", level=2)
    doc.add_paragraph(
        "Không phải địa điểm nào OSM cũng có giờ mở cửa. Nhóm ghi rõ nguồn của từng giá trị "
        "(hours_source): giá trị thật từ OSM, phần thiếu điền bằng heuristic theo loại hình và "
        "đánh dấu rõ ràng — không trộn lẫn để làm đẹp số liệu.")
    img("fig_hours_source.png", 11)
    caption("Giờ mở cửa theo nguồn — chỉ nguồn thật được đưa vào audit")

    doc.add_heading("3.3 Accuracy Audit — tự kiểm chứng", level=2)
    doc.add_paragraph(
        "Nhóm rút ngẫu nhiên 50 địa điểm, một người đối chiếu tay với Google Maps, chấm đúng/sai "
        "cho tọa độ và giờ mở cửa:")
    img("fig_audit.png", 11)
    caption("Accuracy audit trên 50 địa điểm ngẫu nhiên")
    note(
        "Nhận xét: tọa độ 64% và giờ 48% nghe không cao, nhưng đây chính là giá trị của đồ án — "
        "có công cụ đo được chất lượng dữ liệu mở, và trung thực chỉ ra giới hạn của nó. Ngoài ra "
        "còn có bản đồ độ phủ tương tác reports/coverage_map.html hiển thị cả 500 điểm.")

    # ---------------------------------------------------------------- 4
    doc.add_heading("4. Nhánh 2 — Dữ liệu cảm xúc", level=1)

    doc.add_heading("4.1 Gán nhãn tập gold", level=2)
    doc.add_paragraph(
        "Hai người gán nhãn 493 review tiếng Việt thật vào 3 lớp POS / NEG / NEU. Nhãn chốt "
        "(label) lấy theo annotator2 (người gán chính; annotator1 là vòng rà soát).")
    img("fig_label_dist.png", 11)
    caption("Phân bố 493 nhãn chốt — lệch mạnh về POS")
    note(
        "Dữ liệu lệch mạnh về POS (319/493). Đây là nguyên nhân trực tiếp khiến lớp NEU khó — "
        "mọi model đều ít mẫu NEU để học.")

    doc.add_heading("4.2 Độ đồng thuận — Cohen's Kappa", level=2)
    p = doc.add_paragraph()
    p.add_run("Cohen's Kappa = 0.611").bold = True
    p.add_run(" · đồng ý 78.8% · mức substantial (tốt) theo thang Landis–Koch. "
              "Kappa cao chứng minh tập gold đáng tin để làm chuẩn đánh giá model.")

    doc.add_heading("4.3 So sánh model cảm xúc", level=2)
    doc.add_paragraph("Chạy hai model tiếng Việt có sẵn trên HuggingFace lên tập gold:")
    img("fig_model_compare.png", 13)
    caption("So sánh accuracy & macro-F1")
    table(
        ["Model", "Accuracy", "Macro-F1"],
        [["ViSoBERT (5CD-AI/Vietnamese-Sentiment-visobert)", "86.4%", "0.723"],
         ["PhoBERT (wonrax/phobert-base-vietnamese-sentiment)", "84.2%", "0.656"]],
        widths=[10.5, 3.0, 3.0],
    )
    doc.add_paragraph(
        "Kết luận: chọn ViSoBERT — thắng ở macro-F1, và không cần tách từ nên giữ nguyên "
        "teencode/emoji vốn hay xuất hiện trong review.")

    doc.add_heading("4.4 Ma trận nhầm lẫn", level=2)
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, fname, title in ((t.rows[0].cells[0], "confusion_visobert.png", "ViSoBERT"),
                               (t.rows[0].cells[1], "confusion_phobert.png", "PhoBERT")):
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.add_run().add_picture(str(FIG / fname), width=Cm(7.5))
    doc.add_paragraph()
    note(
        "Nhận xét: cả hai đọc tốt POS và NEG nhưng yếu ở NEU (ViSoBERT F1≈0.40, PhoBERT≈0.26) — "
        "đúng như dự đoán từ dữ liệu lệch ở mục 4.1. Đây là hạn chế nhóm hiểu rõ, không giấu.")

    # ---------------------------------------------------------------- 5
    doc.add_heading("5. Sản phẩm demo — App Streamlit", level=1)
    doc.add_paragraph("Toàn bộ bằng chứng được gói vào một app 3 tab (streamlit run app.py):")
    for line in (
        "Tab 1 — Dữ liệu địa điểm & Chất lượng: tổng quan 500 POI, giờ theo nguồn, accuracy "
        "audit, bản đồ độ phủ, ô tra cứu địa điểm.",
        "Tab 2 — Gán nhãn & So sánh model: phân bố nhãn, Cohen's Kappa, bảng so sánh model, ma "
        "trận nhầm lẫn.",
        "Tab 3 — Demo cảm xúc: nhập nhiều review → model chấm nhãn từng câu + tỷ lệ pos/neg/neu.",
    ):
        doc.add_paragraph(line, style="List Bullet")

    # ---------------------------------------------------------------- 6
    doc.add_heading("6. Hạn chế & hướng phát triển", level=1)
    table(
        ["Hạn chế (trung thực)", "Hướng khắc phục"],
        [["Giờ mở cửa thật chỉ phủ 17% (OSM thiếu)",
          "Bổ sung Google Places (đã có code, hiện bị chặn cấp tổ chức)"],
         ["Tọa độ/giờ audit 64%/48%", "Lọc thêm bằng nhiều nguồn, cập nhật định kỳ"],
         ["NEU yếu do dữ liệu lệch (chỉ 40 mẫu)", "Gán thêm review NEU để cân bằng tập gold"],
         ["Chưa có tính năng lập lộ trình", "Phase 3 (tùy chọn): thuật toán TSP có khung giờ"]],
        widths=[7.5, 8.0],
    )

    # ---------------------------------------------------------------- 7
    doc.add_heading("7. Kết luận", level=1)
    doc.add_paragraph(
        "Đồ án chứng minh nhóm làm chủ quy trình dữ liệu có kiểm chứng cho cả địa điểm và cảm "
        "xúc: thu thập dữ liệu thật → đo được chất lượng (audit 64%/48%) → gán nhãn đáng tin "
        "(Kappa 0.611) → chọn model bằng bằng chứng (ViSoBERT 86.4%) → đóng gói thành app trình "
        "diễn. Điểm mạnh xuyên suốt là tính minh bạch và định lượng — nêu rõ cả điểm mạnh lẫn "
        "giới hạn.")

    note("Phụ lục — tái tạo mọi biểu đồ: python scripts/make_report_figures.py. "
         "Xuất lại file Word này: python scripts/export_docx.py.")

    doc.save(str(OUT))
    print(f"✅ Đã xuất: {OUT.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
