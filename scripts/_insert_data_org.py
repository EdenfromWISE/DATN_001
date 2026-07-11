"""Chèn mục "Cách tổ chức dữ liệu" vào BAO_CAO_DEMO.docx đang có (giữ nguyên phần đã sửa tay).

Đặt mục mới ngay trước "Nhánh 1", đánh số lại các mục sau cho nhất quán.
Chạy 1 lần:  python scripts/_insert_data_org.py
"""
import _bootstrap  # noqa: F401

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor

DOCX = _bootstrap.ROOT / "BAO_CAO_DEMO.docx"
GREY = RGBColor(0x55, 0x55, 0x55)

# Schema bảng pois (SQLite) — lấy đúng theo src/wisetravel/db.py.
POIS_COLS = [
    ("id", "TEXT (khóa chính)", "Mã địa điểm từ OSM, vd node/123, way/456"),
    ("name", "TEXT", "Tên địa điểm"),
    ("lat, lng", "REAL", "Tọa độ vĩ độ / kinh độ"),
    ("category", "TEXT", "food | cafe | lodging | attraction"),
    ("subtype", "TEXT", "Tag OSM gốc, vd amenity=restaurant"),
    ("district", "TEXT", "Quận/phường từ tag addr:* (nếu có)"),
    ("address", "TEXT", "Địa chỉ ghép từ tag addr:* của OSM"),
    ("address_google", "TEXT", "Địa chỉ chuẩn từ Google (nếu làm giàu; hiện trống)"),
    ("opening_hours", "TEXT", "Chuỗi giờ mở cửa kiểu OSM, vd Mo-Su 08:00-22:00"),
    ("hours_source", "TEXT", "Nguồn của giờ: osm | google | heuristic | manual"),
    ("place_id", "TEXT", "Google Place ID (nếu làm giàu)"),
    ("business_status", "TEXT", "OPERATIONAL | CLOSED_TEMPORARILY | CLOSED_PERMANENTLY"),
    ("price_level", "INTEGER", "Mức giá 1..3"),
    ("price_level_estimated", "INTEGER", "1 = ước lượng heuristic, 0 = suy từ dữ liệu thật"),
    ("est_duration_min", "INTEGER", "Thời lượng tham quan ước lượng (phút)"),
    ("source", "TEXT", "Nguồn dữ liệu, vd OpenStreetMap"),
    ("last_updated", "TEXT", "Ngày cập nhật (ISO)"),
]

GOLD_COLS = [
    ("id", "Số thứ tự review"),
    ("text", "Nội dung review tiếng Việt"),
    ("label", "Nhãn CHỐT dùng làm chuẩn (= annotator2): POS | NEG | NEU"),
    ("annotator1", "Nhãn của người gán 1 (vòng rà soát)"),
    ("annotator2", "Nhãn của người gán 2 (người gán chính)"),
    ("source", "Nguồn review"),
]

FILE_ORG = [
    ("data/wisetravel.db", "SQLite — bảng pois (500 địa điểm)"),
    ("data/gold/gold.csv", "Tập gold cảm xúc (493 review đã gán)"),
    ("data/audit_sample.csv", "50 địa điểm đã kiểm tay (coord_dung / hours_dung)"),
    ("reports/", "Kết quả sinh tự động: báo cáo, biểu đồ, ma trận nhầm lẫn"),
]

# Đánh số lại các heading hiện có (khớp theo từ khóa, tránh phụ thuộc dấu câu).
RENUMBER = [
    ("Nhánh 1", "3. Nhánh 1 — Dữ liệu địa điểm"),
    ("Thu thập", "3.1 Thu thập & chuẩn hóa"),
    ("Minh bạch", "3.2 Minh bạch nguồn giờ mở cửa"),
    ("Accuracy Audit", "3.3 Accuracy Audit — tự kiểm chứng"),
    ("Nhánh 2", "4. Nhánh 2 — Dữ liệu cảm xúc"),
    ("Gán nhãn tập gold", "4.1 Gán nhãn tập gold"),
    ("đồng thuận", "4.2 Độ đồng thuận — Cohen's Kappa"),
    ("So sánh model", "4.3 So sánh model cảm xúc"),
    ("Ma trận", "4.4 Ma trận nhầm lẫn"),
    ("Hạn chế", "5. Hạn chế & hướng phát triển"),
]


def set_text(p, text):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.add_run(text)


def main():
    doc = Document(str(DOCX))

    # Bỏ qua nếu đã chèn rồi.
    if any("Cách tổ chức dữ liệu" in p.text for p in doc.paragraphs):
        print("Đã có mục 'Cách tổ chức dữ liệu' — không chèn lại.")
        return

    # 1) Đánh số lại + tìm anchor (heading Nhánh 1).
    anchor = None
    for p in doc.paragraphs:
        if not p.style.name.startswith("Heading"):
            continue
        for kw, new in RENUMBER:
            if kw in p.text:
                set_text(p, new)
                if kw == "Nhánh 1":
                    anchor = p
                break
    if anchor is None:
        raise SystemExit("Không tìm thấy heading 'Nhánh 1' để chèn trước.")

    # 2) Dựng nội dung mục mới (thêm ở cuối rồi dời lên trước anchor).
    new_elems = []

    def H(text, level):
        h = doc.add_heading(text, level=level)
        new_elems.append(h._p)
        return h

    def P(text, italic=False, grey=False):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.italic = italic
        if grey:
            r.font.color.rgb = GREY
            r.font.size = Pt(10)
        new_elems.append(p._p)
        return p

    def T(headers, rows, widths):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Light Grid Accent 1"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for c, h in zip(t.rows[0].cells, headers):
            run = c.paragraphs[0].add_run(h)
            run.bold = True
        for row in rows:
            cells = t.add_row().cells
            for c, v in zip(cells, row):
                c.text = str(v)
        for col, w in zip(t.columns, widths):
            for c in col.cells:
                c.width = Cm(w)
        new_elems.append(t._tbl)
        sp = doc.add_paragraph()
        new_elems.append(sp._p)

    H("2. Cách tổ chức dữ liệu", 1)
    P("Dữ liệu được tổ chức thành hai kho tách biệt theo hai bài toán, cộng các file kiểm "
      "chứng. Địa điểm lưu trong SQLite (một bảng pois, mỗi dòng một địa điểm, khóa chính là "
      "mã OSM); dữ liệu cảm xúc lưu dạng CSV phẳng để hai người dễ cùng gán nhãn.")

    H("2.1 Địa điểm — bảng pois (SQLite)", 2)
    P("Mỗi dòng là một địa điểm. Cột hours_source đánh dấu rõ nguồn của giờ mở cửa để tách "
      "dữ liệu thật khỏi giá trị heuristic khi đo chất lượng.")
    T(["Trường", "Kiểu", "Ý nghĩa"], POIS_COLS, [4.0, 3.5, 8.0])

    H("2.2 Cảm xúc — tập gold (CSV)", 2)
    P("File data/gold/gold.csv, mỗi dòng một review với nhãn của hai người gán và nhãn chốt.")
    T(["Cột", "Ý nghĩa"], GOLD_COLS, [3.5, 12.0])

    H("2.3 File kiểm chứng & kết quả", 2)
    T(["Đường dẫn", "Nội dung"], FILE_ORG, [5.5, 10.0])
    P("Nguyên tắc: dữ liệu thô + nhãn người gán được giữ trong repo; báo cáo và biểu đồ trong "
      "reports/ đều sinh lại được bằng script, đảm bảo tái lập.", italic=True, grey=True)

    # 3) Dời toàn bộ elem mới lên trước anchor, giữ đúng thứ tự.
    for el in new_elems:
        anchor._p.addprevious(el)

    doc.save(str(DOCX))
    print(f"✅ Đã chèn mục 'Cách tổ chức dữ liệu' vào {DOCX.name} (giữ nguyên phần bạn đã sửa).")


if __name__ == "__main__":
    main()
